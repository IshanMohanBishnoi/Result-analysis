def analyze():
    from docx import Document
    import pymysql as db
    from docx.shared import Inches
    import os
    import shutil
    import matplotlib.pyplot as plt
    doc = Document()
    doc.add_heading('CBSE RESULT ANALYSIS', 0)
    Head = 1
    try:
        con = db.connect(host='localhost', database='RESULT', password='G@v!nd', user='root')
        cur = con.cursor()
    except:
        print('Initialize first')
        return
    os.makedirs('graphs',exist_ok=True)

    def create_table(querry:str,why:str,Head:int = 1):
        nonlocal cur,doc
        cur.execute(querry)
        data = cur.fetchall()
        if not data:
            doc.add_paragraph(f"No data found for: {why}")
            return

        doc.add_heading(why,Head)

        headers = [i[0] for i in cur.description]
        table = doc.add_table(rows=1,cols=len(headers))
        table.style = 'Colorful Grid Accent 1'
        hdr_cells = table.rows[0].cells
        for i,header in enumerate(headers):
            hdr_cells[i].text = str(header)

        
        for row in data:
            row_cells = table.add_row().cells
            for i,cell in enumerate(row):
                row_cells[i].text = str(cell)

    create_table('SELECT * FROM PRESENT_DATA','STUDENTS APPERANCE:')

    querry = "SELECT table_name FROM information_schema.tables WHERE table_schema = 'RESULT' AND table_name NOT IN ('PRESENT_DATA', 'STUDENTS')"
    cur.execute(querry)
    subjects =[i[0] for i in cur.fetchall()] 
    def get_result(query):
            nonlocal cur
            cur.execute(query=query)
            data = cur.fetchall()
            return data
        
    def subject_analysis(subject):
        nonlocal doc 
        doc.add_heading("Analysis for: {} :".format(subject.upper()),1)

        T_st = get_result("SELECT COUNT(ROLL) FROM {}".format(subject))[0][0]
        t_pas = get_result("SELECT ((SELECT COUNT(ROLL) FROM {})-(SELECT COUNT(ROLL) FROM {} WHERE GRADES = 'E')) AS TOTAL_PASSED".format(subject,subject))[0][0]
        t_f = get_result("SELECT COUNT(ROLL) FROM {} WHERE GRADES = 'E'".format(subject))[0][0]
        s_avg = get_result("SELECT SUM(MARKS)/COUNT(ROLL) AS SUB_AVG FROM {}".format(subject))[0][0]

        paragraph = '''Total students:{}\nTotal students passed:{}\nTotal students failed{}\nSubject average:{}'''.format(T_st,t_pas,t_f,s_avg)
        doc.add_paragraph(paragraph)

        multiple_lines_commands = {"Students with Highest marks:":"SELECT I.NAME,II.MARKS FROM STUDENTS AS I JOIN {} AS II ON I.ROLL = II.ROLL WHERE II.MARKS = (SELECT MAX(MARKS) FROM {})".format(subject,subject),
                                    "Grade count:": "SELECT GRADES,COUNT(GRADES) AS NO FROM {} GROUP BY GRADES ORDER BY GRADES".format(subject),
                                    "Grade percentage:":"SELECT GRADES,(COUNT(GRADES)/(SELECT COUNT(ROLL) FROM {})*100) AS GRADE_PERCENT FROM {} GROUP BY GRADES ORDER BY GRADES".format(subject,subject),
                                    "Students who failed:":"SELECT I.NAME AS STUDENTS_WHO_FAILED FROM STUDENTS AS I JOIN {} AS II ON I.ROLL = II.ROLL WHERE II.GRADES = 'E'".format(subject,subject),
                                    "Students with 100 score": "SELECT I.NAME AS STUDENTS_WITH_100_SCORE FROM STUDENTS AS I JOIN {} AS II ON I.ROLL = II.ROLL WHERE II.MARKS = 100".format(subject)
                                }
        for i in multiple_lines_commands:
            create_table(multiple_lines_commands[i],i)

        grade_count = get_result(multiple_lines_commands["Grade count:"])
        grades,counts = [],[]
        for i,j in grade_count:
            grades.append(i)
            counts.append(j)
        plt.xlabel('Grades')
        plt.bar(grades,counts,color=['red','yellow','blue','green','red','yellow','blue','green','red'])
        plt.ylabel('Total number')
        plt.title(subject+" "+'Grades')
        plt.savefig(f'graphs/{subject}.png')
        plt.close()
        doc.add_picture(f'graphs/{subject}.png')

    for it in subjects:
        subject_analysis(it)

    def stream_analysis(stream ):
        nonlocal doc
        doc.add_heading("Stream analysis of {}:".format(stream),1)
        t_st_appeared = get_result("SELECT COUNT(ROLL) FROM STUDENTS GROUP BY STREAM HAVING STREAM = '{}'".format(stream))[0][0]
        t_f = get_result("SELECT COUNT(ROLL) AS FAILED_STUDENTS FROM STUDENTS WHERE STREAM = '{}' AND PERCENTAGE < 33.0".format(stream))[0][0]
        t_pas = get_result("SELECT COUNT(ROLL) AS FAILED_STUDENTS FROM STUDENTS WHERE STREAM = '{}' AND PERCENTAGE >= 33.0".format(stream))[0][0]
        doc.add_paragraph(f"Total students appeared:{t_st_appeared}\nFailed students:{t_f}\nPassed students:{t_pas}")
        create_table("SELECT NAME,PERCENTAGE FROM STUDENTS WHERE STREAM = '{}' ORDER BY PERCENTAGE DESC LIMIT 25".format(stream),'Top 25 merit list')

    streams = ('COM','SCI','HUM')
    for i in streams:
        stream_analysis(i)

    def school_level():
        nonlocal doc
        school_avg = get_result("SELECT SUM(PERCENTAGE)/COUNT(ROLL) AS SCHOOL_AVG FROM STUDENTS")[0][0]
        doc.add_paragraph(f"School average:{school_avg}")
        create_table("SELECT NAME, PERCENTAGE FROM STUDENTS WHERE PERCENTAGE = (SELECT MAX(PERCENTAGE) FROM STUDENTS)","⭐⭐⭐Highest Aggregate⭐⭐⭐")

    school_level()
    shutil.rmtree('graphs')   
    year = input('enter the year of exam;')
    doc.save(f'Analysis{year}.docx')
